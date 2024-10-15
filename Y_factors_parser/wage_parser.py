import pandas as pd
import re
import requests
from bs4 import BeautifulSoup as bs
import docx
import os
import datetime

from date_functions import str_month2digit_month, str_digit2month, reformate_quarterly_date, reformate_date
from help_functions import append_date_rez_file_Y, doc_to_docx, append_date_rez_file_X


def pars_year_by_months(links_1):
    links_1.columns = ['Месяц', 'Ссылка', 'Дополнительная ссылка']
    links_1 = links_1.iloc[::-1].reset_index(drop=True)
    return links_1


def download_document(year, month, url):
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links = pd.DataFrame()
    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        dok_link = link.get('href')
        links = links._append([[branch_name, dok_link]])

    indicator_1 = 'Заработная плата и пенсии'
    indicator_2 = 'Денежные доходы'
    if len(links[links[0] == indicator_1][1]) == 0 and len(links[links[0] == indicator_2][1]) == 0:
        print(f'NO DOCUMENTS {year}_{month}: {indicator_1}, {year}_{month}: {indicator_2}')
    else:
        indicator = [indicator_1, indicator_2][len(links[links[0] == indicator_1][1]) == 0]
        link_to_download = links[links[0] == indicator][1].values[0]
        dok_name_to_download = f'{year}_{month}-Зарплаты.doc'  # 2024_02-2-4-0.doc
        folder = os.getcwd()
        folder = os.path.join(folder, 'temp_data', dok_name_to_download)

        response = requests.get(link_to_download, headers=header)
        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document wage {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def parse_docx_document(path, year, month):
    '''
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    '''
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0

    data_table = [[] for _ in range(len(doc.tables[3].rows))]
    for i, row in enumerate(doc.tables[3].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    comment = data_table.iloc[-1, 0]
    data_table = data_table[data_table.iloc[:, 0].str.contains(f'{reformate_quarterly_date(month)}')]
    if data_table.empty:
        data_table = [[] for _ in range(len(doc.tables[2].rows))]
        for i, row in enumerate(doc.tables[2].rows):
            for cell in row.cells:
                data_table[i].append(cell.text)

        data_table = pd.DataFrame(data_table)
        comment = data_table.iloc[-1, 0]
        data_table = data_table[data_table.iloc[:, 0].str.contains(f'{reformate_quarterly_date(month)}')]

    data_table = data_table[data_table.iloc[:, 0].apply(lambda x: len(x)) < 20]
    data_table = data_table.iloc[-1:]
    data_table.iloc[0, 0] = reformate_date(data_table.iloc[0, 0], year)

    return data_table.iloc[0][0], data_table.iloc[0][4], data_table.iloc[0][1], comment


def parse_docx_document_for_file_X(path, year):
    """Функиця для сбора данных для файла rez_file_X"""
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0

    data_table = [[] for _ in range(len(doc.tables[3].rows))]
    for i, row in enumerate(doc.tables[3].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: x if ('Январь-' not in x) else '')
    data_table = data_table[data_table.iloc[:, 0].str.contains('Январь|Февраль|Март|Апрель|Май|Июнь|Июль|Август|Сентябрь|Октябрь|Ноябрь|Декабрь')]
    if data_table.empty:
        data_table = [[] for _ in range(len(doc.tables[2].rows))]
        for i, row in enumerate(doc.tables[2].rows):
            for cell in row.cells:
                data_table[i].append(cell.text)

        data_table = pd.DataFrame(data_table)
        data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: x if ('Январь-' not in x) else '')
        data_table = data_table[data_table.iloc[:, 0].str.contains(
            'Январь|Февраль|Март|Апрель|Май|Июнь|Июль|Август|Сентябрь|Октябрь|Ноябрь|Декабрь')]
    data_table = data_table[data_table.iloc[:, 0].apply(lambda x: len(x)) < 20]
    data_table = data_table[[0, 1, 4, 5]]
    for i in range(len(data_table)):
        if i <= 11:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year - 1)
        else:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year)

    # Update tez file X
    xlsx_path = 'rez_file_X_v6.xlsx'
    data_xlsx = pd.read_excel(xlsx_path)
    if data_table.values[-1][0] not in list(data_xlsx['date']):
        append_date_rez_file_X()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = 'Средняя номинальная заработная плата'
    name_2 = 'Реальная среднемесячная начисленная заработная плата, в % к соответствующему периоду предыдущего года '
    name_3 = 'Реальная среднемесячная начисленная заработная плата,  в % к предыдущему периоду'
    for j in data_table.values:
        data_xlsx.loc[data_xlsx['date'] == j[0], name_1] = float(j[1].replace(' ', ''))
        data_xlsx.loc[data_xlsx['date'] == j[0], name_2] = float(j[2].replace(',', '.'))
        data_xlsx.loc[data_xlsx['date'] == j[0], name_3] = float(j[3].replace(',', '.'))
    data_xlsx.to_excel(xlsx_path, index=False)


def parse_docx_document_kvartal(path, year):
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0
    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for i, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table = data_table[data_table.iloc[:, 0].str.contains('I квартал|I полугодие|Январь-сентябрь|Год')]
    data_table = data_table.loc[data_table[0] != 'II квартал']
    data_table = data_table.loc[data_table[0] != 'III квартал']
    data_table = data_table[data_table.iloc[:, 0].apply(lambda x: len(x)) < 20]
    data_table = data_table[[0, 3]]
    for i in range(len(data_table)):
        if i < 4:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year - 1)
        else:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year)

    return data_table


def update_rez_file_y(data, xlsx_path='rez_file_Y_v2.xlsx'):
    data_xlsx = pd.read_excel(xlsx_path)
    if list(data.keys())[-1] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = 'Реальная заработная плата'
    name_2 = 'Среднемесячная номинальная начисленная заработная плата, рублей'
    for j in data:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j, name_1] = data[j][0]
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j, name_2] = data[j][1]

    data_xlsx.to_excel(xlsx_path, index=False)


def update_rez_file_y_kvartal(data, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    name = 'Реальные располагаемые денежные доходы'
    for j in data.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name] = float(j[1].replace(',', '.'))

    data_xlsx.to_excel(xlsx_path, index=False)


def check_last_month_in_table(links_list, year, xlsx_path='rez_file_Y_v2.xlsx'):
    data_xlsx = pd.read_excel(xlsx_path)
    data_xlsx = data_xlsx[['Целевой показатель', 'Реальная заработная плата']]
    month = data_xlsx.dropna(subset=['Реальная заработная плата']).iloc[-1]['Целевой показатель']
    if year == month.year + 1 and month.month in [11, 12]:
        return links_list, 0
    elif month.month in [11, 12] and year == month.year:
        return links_list[11:], 0
    else:
        month = str_digit2month(str(month).split('-')[1])
        idx = links_list[links_list['Месяц'] == month].index[0]
        return links_list.loc[idx:], idx


def wage_parser_main(link_dict):
    '''
    Основная функция. Выполняет проверку данных на полноту. Скачивет недостающие
    данные и дополняет ими файл с данными.
    '''
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Реальные располагаемые '
                                                                                           'денежные доходы']).iloc[
                                            -1]['Целевой показатель']).year

    if now - last_year_in_table < 1:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table, now + 1):
            years.append(y)
    new_data = {}
    kvartal_data = pd.DataFrame()
    for year in years:
        links_data = pars_year_by_months(link_dict[year])
        links_data, idx = check_last_month_in_table(links_data, year)
        if links_data.empty:
            continue
        else:
            print('Ссылки получены')
            for month in links_data['Месяц']:
                if month != links_data['Месяц'].iloc[-1]:
                    # Скачиваем файл и экспортируем его в докс
                    URL = list(links_data.iloc[links_data[links_data['Месяц'] == month].index + 1 - idx]['Ссылка'])[0]
                    path_to_docfile = download_document(year, month, URL)
                    path = doc_to_docx(path_to_docfile)

                    # Собираем данные из файла для 'Реальная заработная плата' и
                    # 'Среднемесячная номинальная начисленная заработная плата, рублей'
                    if month == 'Январь' or month == 'Январь-март':
                        date, value_1, value_2, comm = parse_docx_document(path, year=int(year) - 1, month='Январь-декабрь')
                        new_data[date] = [float(value_1.replace(',', '.')), int(value_2.replace(' ', ''))]

                    date, value_1, value_2, comm = parse_docx_document(path, year=year, month=month)

                    # Собираем данные из файла для 'Реальные располагаемые денежные доходы':
                    if month in ['Январь-февраль', 'Январь-май', 'Январь-август', 'Январь-ноябрь']:
                        kvartal_data = parse_docx_document_kvartal(path, year=year)
                        kvartal_data.columns = ['Дата', 'Реальные располагаемые денежные доходы']

                    os.remove(path_to_docfile)
                    new_data[date] = [float(value_1.replace(',', '.')), int(value_2.replace(' ', ''))]

                    # собираем данные из файла для X
                    parse_docx_document_for_file_X(path, year)
                else:
                    break
        new_data = dict(sorted(new_data.items()))
        if len(new_data) != 0:
            update_rez_file_y(new_data, xlsx_path='rez_file_Y_v2.xlsx')
        if not kvartal_data.empty:
            update_rez_file_y_kvartal(kvartal_data, xlsx_path='rez_file_Y_v2.xlsx')



