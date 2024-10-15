import pandas as pd
import re
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import requests
from bs4 import BeautifulSoup as bs
import docx
import os
import time
import datetime

from date_functions import str_month2digit_month, reformate_date, str_digit2month
from help_functions import append_date_rez_file_Y, doc_to_docx, append_date_rez_file_X

pd.set_option('display.max_rows', None)
pd.set_option('display.max_columns', None)


def reformate_link_list(links_1):
    links_1.columns = ['Месяц', 'Ссылка', 'Дополнительная ссылка']
    links_1 = links_1.iloc[::-1].reset_index(drop=True)
    return links_1


def download_document(year, month, url):
    indicator = 'Розничная торговля'
    doc_link = ''
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    session = requests.Session()
    retry = Retry(connect=3, backoff_factor=0.5)
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('https://', adapter)

    response = session.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        if branch_name == indicator:
            doc_link = link.get('href')
            break

    if len(doc_link) == 0:
        print(f'NO DOCUMENTS {year}_{month}: {indicator}')
    else:
        link_to_download = doc_link
        dok_name_to_download = f'{year}_{month}-Товарооборот.doc'
        folder = os.getcwd()
        folder = os.path.join(folder, 'temp_data', dok_name_to_download)

        response = session.get(link_to_download, headers=header)

        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document trade turnover {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def parse_docx_document(path, year):
    """
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    """
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0
    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for _, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[_].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: ' ' + str(x))
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: x if ('Январь-' not in x) else '')
    data_table = data_table[data_table.iloc[:, 0].str.contains('Январь|Февраль|Март|Апрель|Май|Июнь|Июль|Август|Сентябрь|Октябрь|Ноябрь|Декабрь')]

    for i in [1, 2, ]:
        data_table.iloc[:, i] = data_table.iloc[:, i].str.replace(' ', '').str.replace('\xa0', '').str.replace(',', '.')
    data_table = data_table[[0, 3, 5, 7]]

    for i in range(len(data_table)):
        if i <= 11:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year - 1)
        else:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year)

    print(f'Document trade turnover {path} was parsed')
    return data_table[[0, 3]], data_table[[0, 5]], data_table[[0, 7]]


def parse_docx_document_kvartal(path, year):
    '''
    Функция осуществляет парсинг документа для квартальных данных
    '''
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0
    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for i, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table = data_table[data_table.iloc[:, 0].str.contains('I квартал|I полугодие|Январь-сентябрь|Год')]
    data_table = data_table.loc[data_table[0] != 'II квартал']
    data_table = data_table.loc[data_table[0] != 'III квартал']
    data_table = data_table[data_table.iloc[:, 0].apply(lambda x: len(x)) < 20]
    data_table = data_table[[0, 5]]
    for i in range(len(data_table)):
        if i < 4:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year - 1)
        else:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year)
    return data_table


def update_rez_file_y(data, kvartal_data, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    if list(data.keys())[-1] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = 'Розничный товарооборот'
    name_2 = 'Розничный товарооборот, темп роста, % г/г'
    for j in data.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_1] = float(
            str(j[1]).replace(',', '.').replace(' ', ''))
    if len(kvartal_data) != 0:
        for c in kvartal_data.values:
            data_xlsx.loc[data_xlsx['Целевой показатель'] == c[0], name_2] = float(
                str(c[1]).replace(',', '.'))

    data_xlsx.to_excel(xlsx_path, index=False)


def update_rez_file_X(data, column_name, xlsx_path='rez_file_X_v6.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_X_v6.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    data.columns = ['date', 'value']
    if list(data['date'])[-1] not in list(data_xlsx['date']):
        append_date_rez_file_X()
        data_xlsx = pd.read_excel(xlsx_path)

    for j in list(data.values):
        data_xlsx.loc[data_xlsx['date'] == j[0], column_name] = float(j[1].replace(',', '.').replace(' ', ''))

    data_xlsx.to_excel(xlsx_path, index=False)


def trade_turnover_parser_main(link_dict):
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Розничный товарооборот']).iloc[
                                            -1]['Целевой показатель']).year
    kvartal_data = {}

    if now - last_year_in_table < 1:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table, now + 1):
            years.append(y)

    for year in years:
        links_data = reformate_link_list(link_dict[year])
        if links_data.empty:
            continue
        else:
            month = links_data['Месяц'].iloc[-1]
            URL = links_data['Ссылка'].iloc[-1]
            print(month, URL)
            time.sleep(5)
            path_to_docfile = download_document(year, month, URL)
            print(path_to_docfile)
            path = doc_to_docx(path_to_docfile)

            new_data_for_y, data_for_x_1, data_for_x_2 = parse_docx_document(path, year=year)

            if month in ['Январь-март', 'Январь-июнь', 'Январь-сентябрь', 'Январь-декабрь']:
                kvartal_data = parse_docx_document_kvartal(path, year=year)

            os.remove(path=path_to_docfile)

            if len(new_data_for_y) != 0:
                # Обновляем файл Y
                update_rez_file_y(new_data_for_y, kvartal_data, xlsx_path='rez_file_Y_v2.xlsx')
                # Обновляем файл X
                update_rez_file_X(new_data_for_y, 'ДИНАМИКА ОБОРОТА РОЗНИЧНОЙ ТОРГОВЛИ')
                update_rez_file_X(data_for_x_1, 'ДИНАМИКА ОБОРОТА РОЗНИЧНОЙ ТОРГОВЛИ В % к соответствующему периоду предыдущего года')
                update_rez_file_X(data_for_x_2, 'ДИНАМИКА ОБОРОТА РОЗНИЧНОЙ ТОРГОВЛИ В % к предыдущему периоду')


