import pandas as pd
import re
import requests
from bs4 import BeautifulSoup as bs
import docx
import os
import time
import datetime

from date_functions import str_month2digit_month, reformate_date
from help_functions import append_date_rez_file_Y, doc_to_docx


def pars_year_by_months(links_1):
    # ВВП обновлляется 4 раза в год
    links_1.columns = ['Месяц', 'Ссылка', 'Дополнительная ссылка']
    links_1 = links_1[links_1['Месяц'].str.contains('май|август|ноябрь')]
    links_1 = links_1.iloc[::-1].reset_index(drop=True)

    if links_1.empty:
        return 0, 0
    else:
        return links_1['Месяц'].iloc[-1], links_1['Ссылка'].iloc[-1]


def download_document(year, month, url):
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links = pd.DataFrame()
    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        dok_link = link.get('href')
        links = links._append([[branch_name, dok_link]])

    indicator = 'Объем ВВП'
    if len(links[links[0] == indicator][1]) == 0:
        print(f'NO DOCUMENT {year}_{month}: {indicator}')
    else:
        link_to_download = links[links[0] == indicator][1].values[0]
        dok_name_to_download = f'{year}_{month}-ВВП.doc'
        folder = os.getcwd()
        folder = os.path.join(folder, 'temp_data', dok_name_to_download)

        response = requests.get(link_to_download, headers=header)
        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document DGP {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def parse_docx_document(path, year):
    """
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    """
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0

    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for i, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: ' ' + str(x))
    data_table = data_table[data_table.iloc[:, 0].str.contains('I квартал|I полугодие|Январь-сентябрь|IV квартал|Год')]

    if data_table.empty:
        data_table = [[] for _ in range(len(doc.tables[1].rows))]
        for i, row in enumerate(doc.tables[1].rows):
            for cell in row.cells:
                data_table[i].append(cell.text)

    for i in [1, 2, ]:
        data_table.iloc[:, i] = data_table.iloc[:, i].str.replace(' ', '').str.replace('\xa0', '').str.replace(',', '.')
    print(f'Document DGP {path} was parsed')

    df_1 = data_table[[0, 1]]
    df_2 = data_table[[0, 1]]
    df_1 = df_1[df_1.iloc[:, 0].str.contains(' I квартал|I полугодие|Январь-сентябрь|Год')]
    df_2 = df_2[df_2.iloc[:, 0].str.contains(' I квартал|II квартал|III квартал|IV квартал')]

    for i in range(len(df_1)):
        if i < 4:
            df_1.iloc[i, 0] = reformate_date(df_1.iloc[i, 0], year - 1)
            df_2.iloc[i, 0] = reformate_date(df_2.iloc[i, 0], year - 1)
        else:
            df_1.iloc[i, 0] = reformate_date(df_1.iloc[i, 0], year)
            df_2.iloc[i, 0] = reformate_date(df_2.iloc[i, 0], year)

    return df_1, df_2


def update_rez_file_y(df_1, df_2, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    if df_1.values[-1][0] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = 'Реальный ВВП (Темп роста накопленным итогом)'
    name_2 = 'Реальный ВВП (Темп роста, квартал к кварталу прошлого года)'
    for j in df_1.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_1] = float(j[1].replace(',', '.'))
    for k in df_2.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == k[0], name_2] = float(k[1].replace(',', '.'))

    data_xlsx.to_excel(xlsx_path, index=False)


def DGP_parser_main(link_dict):
    """
    Основная функция. Выполняет проверку данных на полноту. Скачивет недостающие
    данные и дополняет ими файл с данными.
    """
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Реальный ВВП (Темп роста накопленным итогом)']).iloc[
                                            -1]['Целевой показатель']).year
    if now - last_year_in_table < 2:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table + 1, now + 1):
            years.append(y)
    for year in years:
        month, URL = pars_year_by_months(link_dict[year])
        print(month, URL)
        if month == 0:
            print('Новых данных по ВВП не появилось')
            break
        time.sleep(15)
        path_to_docfile = download_document(year, month, URL)
        path = doc_to_docx(path_to_docfile)
        df_1, df_2 = parse_docx_document(path, year=year)
        os.remove(path_to_docfile)
        if not df_1.empty and not df_2.empty:
            update_rez_file_y(df_1, df_2)
