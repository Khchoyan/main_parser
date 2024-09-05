import pandas as pd
import re
import docx
import os
import requests
from bs4 import BeautifulSoup as bs
import datetime

from date_functions import str_month2digit_month, reformate_date
from help_functions import append_date_rez_file_Y, doc_to_docx


def pars_year_by_months(links_1):
    # Инвестиции выходят 4 раза в год
    links_1 = links_1[links_1[0].str.contains('февраль|апрель|июль|октябрь')]
    return links_1.iloc[0, 0].split('-')[-1], links_1.iloc[0, 1]


def download_document(year, month, url):
    '''
    Функция скачивает документ с данными по инвестициям за конкретный месяц.
    year - год в формате ХХХХ.
    month - полное название месяца на русском языке.
    url - ссылка на документ.
    Первые две переменные необходимы для назначения имени скачиваемому файлу.
    Возвращает путь к сохранённому файлу.
    '''
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    response = requests.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    links = pd.DataFrame()
    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        dok_link = link.get('href')
        links = links._append([[branch_name, dok_link]])

    indicator = 'Инвестиции в нефинансовые активы'
    if len(links[links[0] == indicator][1]) == 0:
        print(f'NO DOCUMENT {year}_{month}: {indicator}')
    else:
        link_to_download = links[links[0] == indicator][1].values[0]
        if month[0] == '0':
            month = '12' if month[1] == 1 else '0' + str(int(month) - 1)
        else:
            month = '09' if month[1] == 0 else '1' + str(int(month) - 1)

        dok_name_to_download = f'{year}_{month}-Инвестиции.doc'
        folder = os.getcwd()
        folder = os.path.join(folder, 'temp_data', dok_name_to_download)

        response = requests.get(link_to_download, headers=header)
        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document invest {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def parse_docx_document(path, year, month):
    '''
    Функция осуществляет парсинг документа.
    path - путь к документу (обязательно в формате .docx)
    year - текущий год
    '''
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')

    data_table = [[] for _ in range(len(doc.tables[1].rows))]
    for i, row in enumerate(doc.tables[1].rows):
        for cell in row.cells:
            data_table[i].append(cell.text)

    data_table = pd.DataFrame(data_table)
    comment = data_table.iloc[-1, 0]

    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: ' ' + str(x))
    data_table = data_table[data_table.iloc[:, 0].str.contains(' I квартал|I полугодие|Январь-сентябрь|Год')]
    data_table = data_table[data_table.iloc[:, 0].apply(lambda x: len(x)) < 20]
    for i in range(len(data_table)):
        if ')' in data_table.iloc[i, 0]:
            data_table.iloc[i, 0] = data_table.iloc[i, 0][:-2]

    if month == 'февраль':
        year -= 1
    for i in range(len(data_table)):
        if i < 4:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year - 1)
        else:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year)

    for i in [1, 2, ]:
        data_table.iloc[:, i] = data_table.iloc[:, i].str.replace(' ', '').str.replace('\xa0', '').str.replace(',', '.')
        try:
            data_table.iloc[:, i] = data_table.iloc[:, i].astype('float')
        except ValueError:
            print('parse_docx_document: Could not convert string to float. Unknown symbol.')
    print(f'Document invest {path} was parsed')

    return data_table, comment


def update_rez_file_y(data, xlsx_path='rez_file_Y_v2.xlsx'):
    """
        Функция осуществляет обновление файла со всеми данными rez_file_Y_v2.xlsx
    """
    data_xlsx = pd.read_excel(xlsx_path)
    if data.values[-1][0] not in list(data_xlsx['Целевой показатель']):
        append_date_rez_file_Y()
        data_xlsx = pd.read_excel(xlsx_path)
    name_1 = data.columns[1]
    name_2 = data.columns[2]
    for j in data.values:
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_1] = j[1]
        data_xlsx.loc[data_xlsx['Целевой показатель'] == j[0], name_2] = j[2]

    data_xlsx.to_excel(xlsx_path, index=False)


def run_invest_main(link_dict):
    '''
    Основная функция. Выполняет проверку данных на полноту. Скачивет недостающие
    данные и дополняет ими файл с данными.
    '''
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Инвестиции в основной '
                                                                                           'капитал накопленным '
                                                                                           'итогом, млрд руб']).iloc[
                                            -1]['Целевой показатель']).year

    if now - last_year_in_table < 2:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table + 1, now + 1):
            years.append(y)

    new_data = pd.DataFrame()
    for year in years:
        month, url = pars_year_by_months(link_dict[year])
        print(month, url)

        path_to_docfile = download_document(year, month, url)

        path = doc_to_docx(path_to_docfile)
        df, comm = parse_docx_document(path, year=year, month=month)
        os.remove(path)

        temp = df.iloc[:, :3]
        temp['3'] = comm
        temp.columns = ['Целевой показатель',
                        'Инвестиции в основной капитал накопленным итогом, млрд руб',
                        'Инвестиции, % накопленным итогом год к году',
                        'Комментарий']
        new_data = new_data._append(temp)

    new_data = new_data.drop_duplicates(subset=['Целевой показатель'], keep='last')
    del new_data['Комментарий']
    update_rez_file_y(new_data, xlsx_path='rez_file_Y_v2.xlsx')
