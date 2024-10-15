import pandas as pd
import re
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import requests
from bs4 import BeautifulSoup as bs
import docx
import os
import time
import datetime

from date_functions import str_month2digit_month, reformate_date
from help_functions import doc_to_docx, append_date_rez_file_X

pd.set_option('display.max_rows', None)
pd.set_option('display.max_columns', None)


def reformate_link_list(links_1):
    links_1.columns = ['Месяц', 'Ссылка', 'Дополнительная ссылка']
    links_1 = links_1.iloc[::-1].reset_index(drop=True)
    return links_1


def download_document(year, month, url, indicator):
    doc_link = ''
    header = {
        'user-agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:86.0) Gecko/20100101 Firefox/86.0'
    }
    month = str_month2digit_month(month)
    session = requests.Session()
    retry = Retry(connect=3, backoff_factor=0.5)
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('https://', adapter)

    response = session.get(url, headers=header)
    soup = bs(response.content, "html.parser")

    for link in soup.find_all('a'):
        branch_name = link.text
        branch_name = branch_name.replace('\n', '').replace('\r', '').strip()
        branch_name = re.sub(' +', ' ', branch_name)
        if branch_name == indicator:
            doc_link = link.get('href')
            break

    if len(doc_link) == 0:
        print(f'NO DOCUMENTS {year}_{month}: {indicator}')
    else:
        link_to_download = doc_link
        dok_name_to_download = f'{year}_{month}-{indicator}.doc'
        folder = os.getcwd()
        folder = os.path.join(folder, 'temp_data', dok_name_to_download)

        response = session.get(link_to_download, headers=header)

        if response.status_code == 200:
            with open(folder, 'wb') as f:
                f.write(response.content)
            print(f'Document {indicator} {year}_{month} was downloaded.')
        else:
            print('FAILED:', link_to_download)

        return folder


def parse_docx_document(path, year):
    try:
        doc = docx.Document(path)
    except:
        print('parse_docx_document: It is not word document')
        return 0, 0, 0
    data_table = [[] for _ in range(len(doc.tables[0].rows))]
    for _, row in enumerate(doc.tables[0].rows):
        for cell in row.cells:
            data_table[_].append(cell.text)

    data_table = pd.DataFrame(data_table)
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: ' ' + str(x))
    data_table.iloc[:, 0] = data_table.iloc[:, 0].apply(lambda x: x if ('Январь-' not in x) else '')
    data_table = data_table[data_table.iloc[:, 0].str.contains('Январь|Февраль|Март|Апрель|Май|Июнь|Июль|Август|Сентябрь|Октябрь|Ноябрь|Декабрь')]

    for i in [1, 2, ]:
        data_table.iloc[:, i] = data_table.iloc[:, i].str.replace(' ', '').str.replace('\xa0', '').str.replace(',', '.')

    for i in range(len(data_table)):
        if i <= 11:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year - 1)
        else:
            data_table.iloc[i, 0] = reformate_date(data_table.iloc[i, 0], year)

    return data_table


def update_rez_file_X(data, column_name, xlsx_path='rez_file_X_v6.xlsx'):
    data_xlsx = pd.read_excel(xlsx_path)
    data.columns = ['date', 'value_1', 'value_2', 'value_3']
    if list(data['date'])[-1] not in list(data_xlsx['date']):
        append_date_rez_file_X()
        data_xlsx = pd.read_excel(xlsx_path)

    for j in list(data.values):
        data_xlsx.loc[data_xlsx['date'] == j[0], column_name[0]] = float(j[1].replace(',', '.').replace(' ', ''))
        data_xlsx.loc[data_xlsx['date'] == j[0], column_name[1]] = float(j[2].replace(',', '.').replace(' ', ''))
        data_xlsx.loc[data_xlsx['date'] == j[0], column_name[2]] = float(j[3].replace(',', '.').replace(' ', ''))

    data_xlsx.to_excel(xlsx_path, index=False)


def update_rez_file_X_for_agriculture(data, column_name, xlsx_path='rez_file_X_v6.xlsx'):
    data_xlsx = pd.read_excel(xlsx_path)
    data.columns = ['date', 'value_1', 'value_2']
    pattern = r"([0-9]+,[0-9]+)"
    for j in list(data.values):
        if 'в' in j[1]:
            m = re.search(pattern, j[1])
            if m:
                data_xlsx.loc[data_xlsx['date'] == j[0], column_name[0]] = float(m.group(1).replace(',', '.').replace(' ', '')) * 100
            else:
                m = re.search(r"([0-9]+.[0-9]+)", j[1])
                data_xlsx.loc[data_xlsx['date'] == j[0], column_name[0]] = float(
                    m.group(1).replace(',', '.').replace(' ', '')) * 100
        if 'в' in j[2]:
            m = re.search(pattern, j[2])
            if m:
                data_xlsx.loc[data_xlsx['date'] == j[0], column_name[0]] = float(m.group(1).replace(',', '.').replace(' ', '')) * 100
            else:
                m = re.search(r"([0-9]+.[0-9]+)", j[2])
                data_xlsx.loc[data_xlsx['date'] == j[0], column_name[0]] = float(
                    m.group(1).replace(',', '.').replace(' ', '')) * 100
        else:
            data_xlsx.loc[data_xlsx['date'] == j[0], column_name[0]] = float(j[1].replace(',', '.').replace(' ', ''))
            data_xlsx.loc[data_xlsx['date'] == j[0], column_name[1]] = float(j[2].replace(',', '.').replace(' ', ''))
    data_xlsx.to_excel(xlsx_path, index=False)


def trade_turnover_dynamics_of_indicators_parser_main(link_dict):
    now = datetime.datetime.now().year
    last_year_in_table = pd.to_datetime(pd.read_excel('rez_file_Y_v2.xlsx').dropna(subset=['Розничный товарооборот']).iloc[
                                            -1]['Целевой показатель']).year

    file_name_lst = ['Строительство', 'Рестораны, кафе и бары', 'Рынок платных услуг населению', 'Оптовая торговля',
                     'Сельское хозяйство']

    if now - last_year_in_table < 1:
        years = [now]
    else:
        years = []
        for y in range(last_year_in_table, now + 1):
            years.append(y)
    for year in years:
        links_data = reformate_link_list(link_dict[year])
        month = links_data['Месяц'].iloc[-1]
        URL = links_data['Ссылка'].iloc[-1]
        print(month, URL)
        for indicator in file_name_lst:
            path_to_docfile = download_document(year, month, URL, indicator)
            print(path_to_docfile)
            path = doc_to_docx(path_to_docfile)
            data = parse_docx_document(path, year=year)
            time.sleep(5)
            os.remove(path=path_to_docfile)
            if indicator == 'Сельское хозяйство':
                update_rez_file_X_for_agriculture(data, ['Динамика производства продукции сельского хозяйства в % к соответствующему периоду предыдущего года',
                                                       'Динамика производства продукции сельского хозяйства в % к предыдущему периоду'])

            elif indicator == 'Строительство':
                update_rez_file_X(data, ['ДИНАМИКА ОБЪЕМА СТРОИТЕЛЬСТВА',
                                         'ДИНАМИКА ОБЪЕМА СТРОИТЕЛЬСТВА в % к соответствующему периоду  предыдущего года',
                                         'ДИНАМИКА ОБЪЕМА СТРОИТЕЛЬСТВА в % к предыдущему периоду'])
            elif indicator == 'Рестораны, кафе и бары':
                update_rez_file_X(data[[0, 3, 5, 7]], ['ДИНАМИКА ОБОРОТА ОБЩЕСТВЕННОГО ПИТАНИЯ',
                                                       'ДИНАМИКА ОБОРОТА ОБЩЕСТВЕННОГО ПИТАНИЯ в % к соответствующему периоду  предыдущего года',
                                                       'ДИНАМИКА ОБОРОТА ОБЩЕСТВЕННОГО ПИТАНИЯ в % к предыдущему периоду'])
            elif indicator == 'Рынок платных услуг населению':
                update_rez_file_X(data, ['ДИНАМИКА ОБЪЕМА ПЛАТНЫХ УСЛУГ НАСЕЛЕНИЮ',
                                         'ДИНАМИКА ОБЪЕМА ПЛАТНЫХ УСЛУГ НАСЕЛЕНИЮ в % к соответствующему периоду  предыдущего года',
                                         'ДИНАМИКА ОБЪЕМА ПЛАТНЫХ УСЛУГ НАСЕЛЕНИЮ в % к предыдущему периоду'])
            elif indicator == 'Оптовая торговля':
                update_rez_file_X(data[[0, 1, 2, 3]], ['ДИНАМИКА ОБОРОТА ОПТОВОЙ ТОРГОВЛИ',
                                                     'ДИНАМИКА ОБОРОТА ОПТОВОЙ ТОРГОВЛИ в % к соответствующему периоду  предыдущего года',
                                                     'ДИНАМИКА ОБОРОТА ОПТОВОЙ ТОРГОВЛИ в % к предыдущему периоду'])
